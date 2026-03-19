"""
Export module for generating reports in various formats
"""
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.colors import HexColor
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.units import inch
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import csv
from datetime import datetime
from typing import List, Optional
import io

class ExportManager:
    """Manage exports in multiple formats"""
    
    @staticmethod
    def to_pdf(data: List[dict], filename: str, title: str = "Report") -> str:
        """Export data to PDF"""
        output = filename
        doc = SimpleDocTemplate(output, pagesize=letter)
        elements = []
        
        # Title
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=HexColor('#1F4788'),
            spaceAfter=30,
            alignment=1
        )
        elements.append(Paragraph(title, title_style))
        elements.append(Spacer(1, 0.3*inch))
        
        # Data table
        if data:
            headers = list(data[0].keys())
            table_data = [headers]
            for row in data:
                table_data.append([str(row.get(h, '')) for h in headers])
            
            table = Table(table_data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), HexColor('#1F4788')),
                ('TEXTCOLOR', (0, 0), (-1, 0), '#FFFFFF'),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('GRID', (0, 0), (-1, -1), 1, '#CCCCCC'),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), ['#FFFFFF', '#F0F0F0']),
            ]))
            elements.append(table)
        
        # Timestamp
        elements.append(Spacer(1, 0.3*inch))
        timestamp = datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S")
        elements.append(Paragraph(f"Generated on {timestamp}", styles['Normal']))
        
        doc.build(elements)
        return output
    
    @staticmethod
    def to_excel(data: List[dict], filename: str) -> str:
        """Export data to Excel"""
        df = pd.DataFrame(data)
        df.to_excel(filename, index=False)
        return filename
    
    @staticmethod
    def to_csv(data: List[dict], filename: str) -> str:
        """Export data to CSV"""
        df = pd.DataFrame(data)
        df.to_csv(filename, index=False)
        return filename
    
    @staticmethod
    def to_word(data: List[dict], filename: str, title: str = "Report") -> str:
        """Export data to Word document"""
        doc = Document()
        
        # Title
        title_paragraph = doc.add_heading(title, 0)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Timestamp
        timestamp = datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S")
        doc.add_paragraph(f"Generated on {timestamp}")
        
        # Table
        if data:
            headers = list(data[0].keys())
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Light Grid Accent 1'
            
            # Header row
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = str(header)
                # Make header bold
                for paragraph in hdr_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Data rows
            for row_data in data:
                row_cells = table.add_row().cells
                for i, header in enumerate(headers):
                    row_cells[i].text = str(row_data.get(header, ''))
        
        doc.save(filename)
        return filename
    
    @staticmethod
    def to_json(data: List[dict], filename: str) -> str:
        """Export data to JSON"""
        import json
        with open(filename, 'w') as f:
            json.dump(data, f, indent=2, default=str)
        return filename

class RankingExporter:
    """Export ranking data in various formats"""
    
    def __init__(self, rankings: List[dict]):
        self.rankings = rankings
    
    def format_for_export(self) -> List[dict]:
        """Format ranking data for export"""
        formatted = []
        for rank in self.rankings:
            formatted.append({
                'Rank': rank.get('rank_position'),
                'Student Name': rank.get('student_name'),
                'Email': rank.get('student_email'),
                'Department': rank.get('department'),
                'Event': rank.get('event'),
                'Year': rank.get('academic_year'),
                'Marks': rank.get('marks'),
                'Grade': rank.get('grade')
            })
        return formatted
    
    def export_pdf(self, filename: str) -> str:
        """Export rankings to PDF"""
        return ExportManager.to_pdf(
            self.format_for_export(),
            filename,
            title="Merit Rankings"
        )
    
    def export_excel(self, filename: str) -> str:
        """Export rankings to Excel"""
        return ExportManager.to_excel(
            self.format_for_export(),
            filename
        )
    
    def export_csv(self, filename: str) -> str:
        """Export rankings to CSV"""
        return ExportManager.to_csv(
            self.format_for_export(),
            filename
        )
    
    def export_word(self, filename: str) -> str:
        """Export rankings to Word"""
        return ExportManager.to_word(
            self.format_for_export(),
            filename,
            title="Merit Rankings Report"
        )

class CertificateReportExporter:
    """Export certificate reports"""
    
    def __init__(self, certificates: List[dict]):
        self.certificates = certificates
    
    def format_for_export(self) -> List[dict]:
        """Format certificate data for export"""
        formatted = []
        for cert in self.certificates:
            formatted.append({
                'Certificate #': cert.get('certificate_number'),
                'Student': cert.get('student_name'),
                'Email': cert.get('student_email'),
                'Course': cert.get('course'),
                'Marks': cert.get('marks'),
                'Status': cert.get('status'),
                'Issue Date': cert.get('issue_date'),
                'Generated': cert.get('created_at')
            })
        return formatted
    
    def export_pdf(self, filename: str) -> str:
        """Export to PDF"""
        return ExportManager.to_pdf(
            self.format_for_export(),
            filename,
            title="Certificate Report"
        )
    
    def export_excel(self, filename: str) -> str:
        """Export to Excel"""
        return ExportManager.to_excel(
            self.format_for_export(),
            filename
        )
    
    def export_csv(self, filename: str) -> str:
        """Export to CSV"""
        return ExportManager.to_csv(
            self.format_for_export(),
            filename
        )
    
    def export_word(self, filename: str) -> str:
        """Export to Word"""
        return ExportManager.to_word(
            self.format_for_export(),
            filename,
            title="Certificate Report"
        )
